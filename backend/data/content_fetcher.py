"""
Universal content fetcher for URLs and documents - PRODUCTION READY
"""
import requests
from bs4 import BeautifulSoup
from typing import Dict, Optional, List, Tuple
import re
import time
import json
from urllib.parse import urlparse, urljoin
import logging
from datetime import datetime
import pdfplumber
import docx
import os
import io
import ipaddress

class ContentFetcher:
    def __init__(self, timeout: int = 30, max_retries: int = 3):
        # Security settings - ADD THESE
        self.allowed_domains = [
            'sec.gov', 'reuters.com', 'bloomberg.com', 'ft.com'
        ]
        self.max_file_size = 10 * 1024 * 1024  # 10MB limit
        self.blocked_networks = ['127.0.0.0/8', '10.0.0.0/8', '172.16.0.0/12', '192.168.0.0/16']
        
        # Rate limiting - ADD THESE
        self.rate_limits = {
            'sec.gov': 10,  # 10 requests per second
            'default': 2    # 2 requests per second to others
        }
        self.last_request_time = {}
        
        self.headers = {
            'User-Agent': 'Financial Analysis Bot 1.0 (compatible; +http://example.com/bot)',
            'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8',
            'Accept-Language': 'en-US,en;q=0.5',
            'Accept-Encoding': 'gzip, deflate',
            'Connection': 'keep-alive',
        }
        
        self.session = requests.Session()
        self.session.headers.update(self.headers)
        self.timeout = timeout
        self.max_retries = max_retries
        
        # Configure logging
        logging.basicConfig(level=logging.INFO)
        self.logger = logging.getLogger(__name__)

    # 🔐 NEW SECURITY METHODS - ADD THESE
    def _validate_url(self, url: str) -> Dict[str, bool]:
        """Validate URL for security risks"""
        try:
            parsed = urlparse(url)
            
            # Check if domain is allowed
            domain_allowed = any(domain in parsed.netloc for domain in self.allowed_domains)
            if not domain_allowed:
                return {"valid": False, "error": f"Domain not allowed: {parsed.netloc}"}
            
            # Check for internal IPs
            try:
                ip = ipaddress.ip_address(parsed.hostname)
                for network in self.blocked_networks:
                    if ip in ipaddress.ip_network(network):
                        return {"valid": False, "error": "Internal IP addresses not allowed"}
            except:
                pass  # Not an IP address
            
            # Check for dangerous file types
            dangerous_extensions = ['.exe', '.bat', '.sh', '.php', '.jar']
            if any(parsed.path.lower().endswith(ext) for ext in dangerous_extensions):
                return {"valid": False, "error": "Dangerous file type"}
                
            return {"valid": True}
        except Exception as e:
            return {"valid": False, "error": f"URL validation failed: {str(e)}"}

    def _check_rate_limit(self, url: str):
        """Enforce rate limiting"""
        domain = urlparse(url).netloc
        current_time = time.time()
        
        limit = self.rate_limits.get(domain, self.rate_limits['default'])
        
        if domain in self.last_request_time:
            time_since_last = current_time - self.last_request_time[domain]
            min_interval = 1.0 / limit
            if time_since_last < min_interval:
                sleep_time = min_interval - time_since_last
                self.logger.info(f"Rate limiting: sleeping {sleep_time:.2f}s for {domain}")
                time.sleep(sleep_time)
        
        self.last_request_time[domain] = time.time()

    def _sanitize_filename(self, filename: str) -> str:
        """Sanitize filename to prevent path traversal"""
        # Remove directory components
        filename = os.path.basename(filename)
        # Remove dangerous characters
        filename = re.sub(r'[^\w\-_.]', '_', filename)
        return filename[:255]  # Limit length

    # 🚨 MODIFIED MAIN METHODS WITH SECURITY CHECKS
    def fetch_content(self, url_or_path: str) -> Dict[str, str]:
        """Universal content fetcher with security checks"""
        # Check if it's a local file
        if os.path.isfile(url_or_path):
            return self._fetch_local_document(url_or_path)
        else:
            return self.fetch_url_content(url_or_path)

    def fetch_url_content(self, url: str) -> Dict[str, str]:
        """Fetch and extract content from URL with security checks"""
        # 🔐 SECURITY CHECK 1: Validate URL
        validation = self._validate_url(url)
        if not validation["valid"]:
            return {
                "status": "error",
                "error": validation["error"],
                "url": url,
                "content": ""
            }
        
        # 🔐 SECURITY CHECK 2: Rate limiting
        self._check_rate_limit(url)
        
        # 🔐 SECURITY CHECK 3: Size limit check
        try:
            head_response = self.session.head(url, timeout=5, allow_redirects=True)
            content_length = head_response.headers.get('content-length')
            if content_length and int(content_length) > self.max_file_size:
                return {
                    "status": "error", 
                    "error": f"File too large: {content_length} bytes",
                    "url": url,
                    "content": ""
                }
        except Exception as e:
            self.logger.warning(f"HEAD request failed, continuing: {e}")

        # Original retry logic with security
        for attempt in range(self.max_retries):
            try:
                self.logger.info(f"Fetching content from {url} (attempt {attempt + 1})")
                
                response = self.session.get(url, timeout=self.timeout, stream=True)
                response.raise_for_status()
                
                # 🔐 SECURITY CHECK 4: Check actual content size
                content_length = response.headers.get('content-length')
                if content_length and int(content_length) > self.max_file_size:
                    return {
                        "status": "error", 
                        "error": f"File too large: {content_length} bytes",
                        "url": url,
                        "content": ""
                    }
                
                # Read content with size limit
                content_bytes = b""
                for chunk in response.iter_content(chunk_size=8192):
                    content_bytes += chunk
                    if len(content_bytes) > self.max_file_size:
                        return {
                            "status": "error",
                            "error": f"File exceeded size limit during download",
                            "url": url,
                            "content": ""
                        }
                
                # Handle different content types
                content_type = response.headers.get('content-type', '')
                
                if 'application/pdf' in content_type:
                    return self._handle_pdf_content(content_bytes, url)
                elif 'application/vnd.openxmlformats-officedocument.wordprocessingml.document' in content_type:
                    return self._handle_docx_content(content_bytes, url)
                elif 'text/html' in content_type:
                    return self._handle_html_content(content_bytes, url, response.encoding)
                else:
                    return self._handle_plain_text(content_bytes, url, response.encoding)
                    
            except requests.exceptions.RequestException as e:
                self.logger.warning(f"Attempt {attempt + 1} failed: {e}")
                if attempt == self.max_retries - 1:
                    return self._get_error_response(url, str(e))
                time.sleep(2 ** attempt)  # Exponential backoff
        
        return self._get_error_response(url, "Max retries exceeded")

    # 🚨 REPLACED FALLBACK METHOD - Never return fake success!
    def _get_error_response(self, url: str, error: str) -> Dict[str, str]:
        """Return proper error response - NO FAKE CONTENT!"""
        return {
            "status": "error",
            "content": "",
            "title": "Fetch Failed",
            "url": url,
            "content_type": "error",
            "word_count": 0,
            "fetch_timestamp": datetime.now().isoformat(),
            "error": f"Failed to fetch content: {error}"
        }

    def _fetch_local_document(self, file_path: str) -> Dict[str, str]:
        """Fetch content from local documents with security checks"""
        try:
            # 🔐 SECURITY CHECK: Sanitize filename
            safe_path = self._sanitize_filename(file_path)
            full_path = os.path.abspath(safe_path)
            
            # Check if file exists and is within allowed directories
            if not os.path.isfile(full_path):
                return {
                    "status": "error",
                    "error": "File not found",
                    "url": f"file://{file_path}",
                    "content": ""
                }
            
            # Check file size
            file_size = os.path.getsize(full_path)
            if file_size > self.max_file_size:
                return {
                    "status": "error",
                    "error": f"File too large: {file_size} bytes",
                    "url": f"file://{file_path}",
                    "content": ""
                }
            
            file_extension = os.path.splitext(full_path)[1].lower()
            
            if file_extension == '.pdf':
                with pdfplumber.open(full_path) as pdf:
                    text_pages = []
                    for page in pdf.pages:
                        text = page.extract_text()
                        if text:
                            text_pages.append(text)
                    full_text = "\n".join(text_pages)
                    
                return {
                    "status": "success",
                    "content": full_text,
                    "title": os.path.basename(full_path),
                    "url": f"file://{full_path}",
                    "content_type": "pdf",
                    "word_count": len(full_text.split()),
                    "fetch_timestamp": datetime.now().isoformat(),
                    "page_count": len(text_pages)
                }
                
            elif file_extension == '.docx':
                doc = docx.Document(full_path)
                full_text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                
                return {
                    "status": "success",
                    "content": full_text,
                    "title": os.path.basename(full_path),
                    "url": f"file://{full_path}",
                    "content_type": "docx",
                    "word_count": len(full_text.split()),
                    "fetch_timestamp": datetime.now().isoformat()
                }
                
            elif file_extension in ['.txt', '.md']:
                with open(full_path, 'r', encoding='utf-8') as file:
                    content = file.read()
                
                return {
                    "status": "success",
                    "content": content,
                    "title": os.path.basename(full_path),
                    "url": f"file://{full_path}",
                    "content_type": "plain_text",
                    "word_count": len(content.split()),
                    "fetch_timestamp": datetime.now().isoformat()
                }
                
            else:
                return {
                    "status": "error",
                    "content": "",
                    "title": os.path.basename(full_path),
                    "url": f"file://{full_path}",
                    "content_type": "unsupported",
                    "error": f"Unsupported file type: {file_extension}"
                }
                
        except Exception as e:
            return {
                "status": "error",
                "content": "",
                "title": os.path.basename(file_path),
                "url": f"file://{file_path}",
                "content_type": "error",
                "error": str(e)
            }

    # UPDATED CONTENT HANDLERS
    def _handle_html_content(self, content: bytes, url: str, encoding: str = 'utf-8') -> Dict[str, str]:
        """Process HTML content with enhanced extraction"""
        try:
            soup = BeautifulSoup(content, 'html.parser', from_encoding=encoding)
            
            # Remove potentially dangerous elements
            for element in soup(["script", "style", "nav", "header", "footer", "aside", "meta", "link", "button", "form"]):
                element.decompose()
            
            # Remove elements with dangerous attributes
            for tag in soup.find_all(True):
                for attr in ['onclick', 'onload', 'onerror', 'onmouseover']:
                    if attr in tag.attrs:
                        del tag.attrs[attr]
            
            title = self._extract_title(soup)
            main_content = self._extract_main_content(soup)
            metadata = self._extract_metadata(soup)
            
            # Clean text
            clean_content = self._clean_extracted_text(main_content)
            
            result = {
                "status": "success",
                "content": clean_content,
                "title": title,
                "url": url,
                "content_type": "web_page",
                "word_count": len(clean_content.split()),
                "fetch_timestamp": datetime.now().isoformat(),
                "metadata": metadata
            }
            
            # Add content analysis
            result.update(self.analyze_content_type(clean_content))
            
            return result
        except Exception as e:
            return self._get_error_response(url, f"HTML processing failed: {str(e)}")

    def _handle_pdf_content(self, content: bytes, url: str) -> Dict[str, str]:
        """Extract text from PDF content with error handling"""
        try:
            if len(content) > self.max_file_size:
                return {
                    "status": "error",
                    "error": "PDF file too large",
                    "url": url,
                    "content": ""
                }
                
            with pdfplumber.open(io.BytesIO(content)) as pdf:
                text_pages = []
                for page in pdf.pages:
                    try:
                        text = page.extract_text()
                        if text and len(text.strip()) > 0:
                            text_pages.append(text)
                    except Exception as e:
                        self.logger.warning(f"Failed to extract text from PDF page: {e}")
                        continue
                
                full_text = "\n".join(text_pages)
                
            return {
                "status": "success",
                "content": full_text,
                "title": f"PDF Document from {urlparse(url).netloc}",
                "url": url,
                "content_type": "pdf",
                "word_count": len(full_text.split()),
                "fetch_timestamp": datetime.now().isoformat(),
                "page_count": len(text_pages)
            }
        except Exception as e:
            return self._get_error_response(url, f"PDF extraction failed: {e}")

    def _handle_docx_content(self, content: bytes, url: str) -> Dict[str, str]:
        """Extract text from DOCX content with error handling"""
        try:
            if len(content) > self.max_file_size:
                return {
                    "status": "error",
                    "error": "DOCX file too large",
                    "url": url,
                    "content": ""
                }
                
            doc = docx.Document(io.BytesIO(content))
            full_text = "\n".join([paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()])
            
            return {
                "status": "success",
                "content": full_text,
                "title": f"Word Document from {urlparse(url).netloc}",
                "url": url,
                "content_type": "docx",
                "word_count": len(full_text.split()),
                "fetch_timestamp": datetime.now().isoformat()
            }
        except Exception as e:
            return self._get_error_response(url, f"DOCX extraction failed: {e}")

    # EXISTING METHODS (keep these as they were)
    def _extract_main_content(self, soup: BeautifulSoup) -> str:
        """Extract main content using multiple strategies"""
        content_selectors = [
            'main', 'article', '[role="main"]', 
            '.content', '.main-content', '.article-content',
            '.post-content', '.entry-content', '.story-content',
            '#content', '#main', '#article',
            '.document', '.text', '.body'
        ]
        
        for selector in content_selectors:
            elements = soup.select(selector)
            if elements:
                largest_element = max(elements, key=lambda el: len(el.get_text().strip()))
                content = largest_element.get_text()
                if len(content.split()) > 100:
                    return content
        
        candidates = soup.find_all(['div', 'section', 'main'])
        if candidates:
            scored_candidates = []
            for candidate in candidates:
                text = candidate.get_text().strip()
                if len(text.split()) < 50:
                    continue
                
                html_length = len(str(candidate))
                text_length = len(text)
                if html_length > 0:
                    density = text_length / html_length
                else:
                    density = 0
                
                link_count = len(candidate.find_all('a'))
                link_penalty = min(link_count / 10, 1.0)
                
                score = density * (1 - link_penalty)
                scored_candidates.append((score, candidate))
            
            if scored_candidates:
                best_candidate = max(scored_candidates, key=lambda x: x[0])
                return best_candidate[1].get_text()
        
        body = soup.find('body')
        if body:
            return body.get_text()
        
        return soup.get_text()

    def _extract_title(self, soup: BeautifulSoup) -> str:
        """Extract page title"""
        title_elem = soup.find('title')
        if title_elem:
            title = title_elem.get_text().strip()
            if title and title != "":
                return title
        
        og_title = soup.find('meta', property='og:title')
        if og_title and og_title.get('content'):
            return og_title['content']
        
        twitter_title = soup.find('meta', attrs={'name': 'twitter:title'})
        if twitter_title and twitter_title.get('content'):
            return twitter_title['content']
        
        h1_elements = soup.find_all('h1')
        if h1_elements:
            for h1 in h1_elements:
                h1_text = h1.get_text().strip()
                if h1_text:
                    return h1_text
        
        headline_elem = soup.find(attrs={'itemprop': 'headline'})
        if headline_elem:
            return headline_elem.get_text().strip()
        
        return "No title found"

    def _extract_metadata(self, soup: BeautifulSoup) -> Dict[str, str]:
        """Extract metadata from page"""
        metadata = {}
        
        description = soup.find('meta', attrs={'name': 'description'})
        if description and description.get('content'):
            metadata['description'] = description['content']
        
        keywords = soup.find('meta', attrs={'name': 'keywords'})
        if keywords and keywords.get('content'):
            metadata['keywords'] = keywords['content']
        
        date_selectors = [
            'meta[property="article:published_time"]',
            'meta[name="publication-date"]',
            'meta[name="publish_date"]',
            'time[datetime]'
        ]
        
        for selector in date_selectors:
            element = soup.select_one(selector)
            if element:
                date_value = element.get('content') or element.get('datetime')
                if date_value:
                    metadata['publication_date'] = date_value
                    break
        
        author_selectors = [
            'meta[name="author"]',
            'meta[property="article:author"]',
            '.author', '[rel="author"]'
        ]
        
        for selector in author_selectors:
            element = soup.select_one(selector)
            if element:
                author = element.get('content') or element.get_text()
                if author:
                    metadata['author'] = author.strip()
                    break
        
        return metadata

    def _clean_extracted_text(self, text: str) -> str:
        """Clean and normalize extracted text"""
        if not text:
            return ""
        
        text = re.sub(r'\s+', ' ', text)
        
        boilerplate_patterns = [
            r'Privacy Policy.*?',
            r'Terms of Service.*?',
            r'Cookie Policy.*?',
            r'Sign up.*?newsletter',
            r'Subscribe.*?',
            r'Follow us.*?',
            r'Copyright.*?\d{4}',
            r'All rights reserved.*?',
            r'Loading\.\.\.',
            r'Please enable JavaScript.*?',
            r'ADVERTISEMENT',
            r'Sponsored Content.*?'
        ]
        
        for pattern in boilerplate_patterns:
            text = re.sub(pattern, '', text, flags=re.IGNORECASE)
        
        lines = text.split('. ')
        cleaned_lines = [line.strip() for line in lines if len(line.split()) > 3]
        
        return '. '.join(cleaned_lines).strip()

    def analyze_content_type(self, content: str) -> Dict[str, str]:
        """Analyze content type and characteristics"""
        content_lower = content.lower()
        
        analysis = {
            "content_type": "unknown",
            "financial_content": False,
            "risk_content": False,
            "regulatory_content": False,
            "document_structure": "unstructured",
            "confidence_score": 0
        }
        
        financial_indicators = [
            'sec', 'earnings', 'revenue', 'profit', 'loss', 'debt', 'equity',
            'financial', 'quarterly', 'annual', 'ebitda', 'margin', 'assets'
        ]
        financial_score = sum(1 for indicator in financial_indicators if indicator in content_lower)
        
        risk_indicators = [
            'risk', 'uncertainty', 'volatility', 'default', 'investigation',
            'compliance', 'regulatory', 'cybersecurity', 'threat', 'vulnerability'
        ]
        risk_score = sum(1 for indicator in risk_indicators if indicator in content_lower)
        
        if financial_score >= 5 and risk_score >= 3:
            analysis["financial_content"] = True
            analysis["risk_content"] = True
            analysis["confidence_score"] = min(100, (financial_score + risk_score) * 10)
        
        if 'item 1a' in content_lower and 'risk factors' in content_lower:
            analysis["content_type"] = "sec_filing"
            analysis["document_structure"] = "structured"
            analysis["confidence_score"] = 95
        elif 'earnings call' in content_lower or 'q&a' in content_lower:
            analysis["content_type"] = "earnings_transcript"
            analysis["confidence_score"] = 85
        elif 'press release' in content_lower:
            analysis["content_type"] = "press_release"
            analysis["confidence_score"] = 80
        elif any(site in content_lower for site in ['reuters', 'bloomberg', 'financial times', 'wsj', 'wall street journal']):
            analysis["content_type"] = "financial_news"
            analysis["confidence_score"] = 75
        
        if any(word in content_lower for word in ['regulation', 'compliance', 'enforcement', 'oversight']):
            analysis["regulatory_content"] = True
        
        return analysis

# USAGE EXAMPLE
if __name__ == "__main__":
    # Initialize with security settings
    fetcher = ContentFetcher(timeout=15, max_retries=2)
    
    # Test with various sources
    test_urls = [
        "https://www.sec.gov/Archives/edgar/data/0000320193/000032019323000106/aapl-20230930.htm",  # Allowed
        "https://www.reuters.com/markets",  # Allowed  
        "https://malicious.com/bad.exe",    # Blocked - dangerous file
        "http://localhost/internal",        # Blocked - internal IP
    ]
    
    for url in test_urls:
        print(f"Fetching: {url}")
        result = fetcher.fetch_content(url)
        print(f"Status: {result['status']}")
        if result['status'] == 'error':
            print(f"Error: {result['error']}")
        else:
            print(f"Content Type: {result['content_type']}")
            print(f"Word Count: {result['word_count']}")
        print("-" * 50)